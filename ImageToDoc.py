import pytesseract
import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files (x86)\\Tesseract-OCR\\tesseract'


def process(imgraw):
    # imgraw = cv2.imread("..\\img2.tif")
    # imgraw = cv2.imread("8.png")
    img = cv2.cvtColor(imgraw, cv2.COLOR_RGB2GRAY)

    height, width = img.shape
    thres, img = cv2.threshold(img, 127, 255, cv2.THRESH_OTSU)
    # inverting
    blacks = sum([1 for i in range(height) for j in range(width) if img[i, j] == 0])
    total = height * width
    percentage = (blacks / (total * 1.0)) * 100
    print "black = " + str(percentage) + "% "
    if percentage < 40:
        img = cv2.bitwise_not(img)

    # opening
    stEl = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
    img = cv2.morphologyEx(img, cv2.MORPH_OPEN, stEl)

    # closing
    stEl = cv2.getStructuringElement(cv2.MORPH_RECT, (100, 25))
    img = cv2.morphologyEx(img, cv2.MORPH_CLOSE, stEl)

    # dilate
    img = cv2.dilate(img, np.ones((15, 9)), 1)

    cv2.imwrite("boundboxtest.tif", img)

    # finding contours
    img, contours, hierarchy = cv2.findContours(img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    nConture = []
    for i in contours:
        # x, y, w, h = cv2.boundingRect(i)
        # if (w < x+30 , h < y+30):
        #     continue
        if cv2.contourArea(i) >= 60:
            nConture.append(i)
    print len(nConture)

    img = cv2.cvtColor(imgraw, cv2.COLOR_RGB2GRAY)
    i = 0
    nConture.reverse()
    document = Document()
    for cnt in nConture:
        x, y, w, h = cv2.boundingRect(cnt)
        cv2.rectangle(imgraw, (x, y), (x + w, y + h), (255, 15, 0), 1)
        cp = img[y:y + h, x:x + w]
        cv2.imwrite("temp\\" + str(i) + ".png", cp)

        text = pytesseract.image_to_string(cp)
        text.strip()
        text.lstrip()
        print "c# " + str(i)
        print len(text)
        if (len(text) > 0):
            paragraph = document.add_paragraph(text)
            if (w / 2) + x > width / 2 - 30 and (w / 2) + x < width / 2 + 30:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif (w / 2) + x > width / 2 + width / 4:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        else:
            cwidth = 8.25 * (w / (1.0 * width))
            document.add_picture("temp\\" + str(i) + ".png", width=Inches(cwidth))
        i = i + 1

    cv2.imwrite("boundbox.tif", imgraw)
    document.save('demo.docx')
